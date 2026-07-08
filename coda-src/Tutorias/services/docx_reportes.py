import re
import docx
import math
from django.http import HttpResponse
from django.shortcuts import redirect
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Inches


def paragraph_replace_text(paragraph, regex, replace_str):
    while True:
        text = paragraph.text
        match = regex.search(text)
        if not match:
            break

        runs = iter(paragraph.runs)
        start, end = match.start(), match.end()

        for run in runs:
            run_len = len(run.text)
            if start < run_len:
                break
            start, end = start - run_len, end - run_len

        run_text = run.text
        run_len = len(run_text)
        run.text = "%s%s%s" % (run_text[:start], replace_str, run_text[end:])
        end -= run_len

        for run in runs:
            if end <= 0:
                break
            run_text = run.text
            run_len = len(run_text)
            run.text = run_text[end:]
            end -= run_len

    return paragraph


def _build_tutor_display_data(tutor):
    est = ""
    dr = ""
    name = f"{tutor.first_name} {tutor.last_name}"

    if tutor.sexo == "M":
        est = "Estimado"
        dr = "Dr."
    if tutor.sexo == "F":
        est = "Estimada"
        dr = "Dra."
    if tutor.second_last_name:
        name = name + f" {tutor.second_last_name}"

    return {
        'estimado': est,
        'nombre_tutor': f"{dr} {name}",
        'nombre_tutor_mayus': f"{dr} {name}".upper(),
    }


def _replace_reporte_placeholders(documento, tutor, oficio, fecha_emision):
    reg_placeh = re.compile(r'\{.*?\}')
    reg_ofi = re.compile(r'\{no_oficio\}')
    reg_fech = re.compile(r'\{fecha\}')
    reg_tut_min = re.compile(r'\{nombre_tutor\}')
    reg_tut = re.compile(r'\{nombre_mayus_tutor\}')
    reg_est = re.compile(r'\{estimado\}')

    tutor_data = _build_tutor_display_data(tutor)

    for p in documento.paragraphs:
        line = p.text
        result = []
        line_matches = [] if (result := re.findall(reg_placeh, line)) is None else result

        for match in line_matches:
            if re.search(reg_ofi, match):
                paragraph_replace_text(p, reg_ofi, f"{oficio}").text
            if re.match(reg_fech, match):
                paragraph_replace_text(p, reg_fech, f"{fecha_emision}").text
            if re.match(reg_tut, match):
                paragraph_replace_text(p, reg_tut, tutor_data['nombre_tutor_mayus'])
            if re.match(reg_est, match):
                paragraph_replace_text(p, reg_est, tutor_data['estimado'])
            if re.match(reg_tut_min, match):
                paragraph_replace_text(p, reg_tut_min, tutor_data['nombre_tutor'])


# Aquí se filtran tutorías (ej: asistencia, estado)
def _build_report_rows(tutorias, mostrar_col_alumno, mostrar_col_fecha, mostrar_col_hora, mostrar_col_tema, mostrar_col_notas, tema_dict):
    rows = []
    for tutoria in tutorias:
        if not tutoria.asistencia:  # deben haber asistido
            continue
        
        if tutoria.alumno.estado != 1:
            continue

        alumno = tutoria.alumno
        row = []

        if mostrar_col_alumno:
            nombre = f"{alumno.first_name} {alumno.last_name}"
            if alumno.second_last_name:
                nombre += f" {alumno.second_last_name}"
            row.append(nombre)

        if mostrar_col_fecha:
            row.append(tutoria.fecha.strftime('%Y-%m-%d'))

        if mostrar_col_hora:
            row.append(tutoria.fecha.strftime('%H:%M'))

        if mostrar_col_tema:
            tema_cod = tutoria.tema[0]
            tema_nombre = tema_dict.get(tema_cod, tema_cod)
            row.append(tema_nombre)

        if mostrar_col_notas:
            row.append(tutoria.observaciones or '')

        rows.append(row)
    return rows


def _insert_blank_paragraph(documento, current_element, text=""):
    paragraph = documento.add_paragraph(text)
    current_element.addnext(paragraph._element)
    return paragraph._element


def _insert_vertical_spacer(documento, current_element, count=1 ): # espacio entre encabezado y tabla
    for _ in range(count):
        current_element = _insert_blank_paragraph(documento, current_element, "")
    return current_element


def _copy_run_style(src_run, dst_run):
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    if src_run.font is not None:
        dst_run.font.name = src_run.font.name
        dst_run.font.size = src_run.font.size
        dst_run.font.bold = src_run.font.bold
        dst_run.font.italic = src_run.font.italic
        dst_run.font.underline = src_run.font.underline
        if src_run.font.color is not None and src_run.font.color.rgb is not None:
            dst_run.font.color.rgb = RGBColor(
                src_run.font.color.rgb[0], src_run.font.color.rgb[1], src_run.font.color.rgb[2]
            )


def _copy_paragraph_style(src_paragraph, dst_paragraph):
    if src_paragraph is None:
        return
    dst_paragraph.alignment = src_paragraph.alignment
    dst_paragraph.style = src_paragraph.style
    dst_fmt = dst_paragraph.paragraph_format
    src_fmt = src_paragraph.paragraph_format
    dst_fmt.left_indent = src_fmt.left_indent
    dst_fmt.right_indent = src_fmt.right_indent
    dst_fmt.first_line_indent = src_fmt.first_line_indent
    dst_fmt.space_before = src_fmt.space_before
    dst_fmt.space_after = src_fmt.space_after
    dst_fmt.line_spacing = src_fmt.line_spacing
    dst_fmt.keep_together = src_fmt.keep_together
    dst_fmt.keep_with_next = src_fmt.keep_with_next
    dst_fmt.page_break_before = src_fmt.page_break_before
    dst_fmt.widow_control = src_fmt.widow_control


def _find_reference_paragraph(documento, needle):
    for p in documento.paragraphs:
        if needle in p.text:
            return p
    return None


def _insert_styled_paragraph_after(documento, current_element, text, reference_paragraph):
    paragraph = documento.add_paragraph()
    run = paragraph.add_run(text)
    if reference_paragraph is not None:
        _copy_paragraph_style(reference_paragraph, paragraph)
        if reference_paragraph.runs:
            _copy_run_style(reference_paragraph.runs[0], run)
    current_element.addnext(paragraph._element)
    return paragraph._element


# Controla encabezado en páginas 2+
def _insert_short_header(documento, current_element, oficio, fecha_emision, refs):
    current_element = _insert_vertical_spacer(documento, current_element, count=6)  # subir/bajar encabezado cambiando este valor
    current_element = _insert_styled_paragraph_after(documento, current_element, f"Oficio No. {oficio}", refs.get('oficio'))
    current_element = _insert_styled_paragraph_after(documento, current_element, f"Ciudad de México a {fecha_emision}", refs.get('fecha'))
    current_element = _insert_styled_paragraph_after(documento, current_element, "Asunto Tutorados(as) atendidos", refs.get('asunto'))
    current_element = _insert_vertical_spacer(documento, current_element, count=3) # espacio entre encabezado y tabla (3 saltos de linea)
    return current_element


# Controla espacio antes de firma
def _push_signature_down(documento, current_element, filas_ultima, max_filas):
    # Calcula cuántas filas NO se usaron en la última página
    faltantes = max_filas - filas_ultima

    # Si hay mucho espacio vacío (tabla corta)
    # → empujar más la firma hacia abajo
    if faltantes >= 8:
        espacios = 5   # aumentar este valor = firma más abajo

    # Espacio medio
    elif faltantes >= 5:
        espacios = 3   # valor intermedio

    # Poco espacio libre
    elif faltantes >= 2:
        espacios = 2   # separación ligera

    # Tabla casi llena
    else:
        espacios = 1   # mínimo espacio para no pegar firma

    # Inserta líneas vacías (cada unidad = un salto de línea)
    return _insert_vertical_spacer(documento, current_element, count=espacios)

def _set_cell_border(cell, color="000000", size="8"):  # color borde, size="8"  # grosor borde (8 ≈ 1pt)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right'):
        tag = f'w:{edge}'
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), size)       # 8 = 1 punto aprox.
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)


def _apply_black_borders(table):
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell)


def _append_field(run, instruction):
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = f' {instruction} '

    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')

    placeholder = OxmlElement('w:t')
    placeholder.text = '1'

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(placeholder)
    run._r.append(fld_char_end)


def _add_page_number_footer(documento):
    for section in documento.sections:
        section.footer_distance = Inches(0.45)  # mover numeración vertical
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.clear()
        paragraph.alignment = 2  # derecha

        run = paragraph.add_run('				Página ')
        _append_field(run, 'PAGE')
        run2 = paragraph.add_run(' de ')
        _append_field(run2, 'NUMPAGES')

def _distribuir_bloques(rows, columnas_activas):
    total = len(rows)
    num_columnas = len(columnas_activas)

    if total == 0:
        return [[]], 10

    # Define cuántas filas caben por página según columnas.
    # 1-2 columnas: tabla más angosta, caben más filas.
    # 3 columnas: caso intermedio.
    # 4 o más: las filas crecen más, caben menos.
    if num_columnas <= 2:
        max_por_pagina = 15
        min_ultima = 1
    elif num_columnas == 3:
        max_por_pagina = 10
        min_ultima = 6
    else:
        max_por_pagina = 8
        min_ultima = 5

    # Reparto inicial por bloques fijos.
    bloques = [rows[i:i + max_por_pagina] for i in range(0, total, max_por_pagina)]

    # Si la última página queda muy vacía, redistribuir.
    # Ejemplo: 21 filas con max 10 → 10,10,1
    # se convierte en 11,10 para evitar una página casi vacía.
    if len(bloques) > 1 and len(bloques[-1]) < min_ultima:
        total_paginas = len(bloques) - 1
        total_filas = total

        filas_base = total_filas // total_paginas
        extra = total_filas % total_paginas

        bloques = []
        inicio = 0
        for i in range(total_paginas):
            tam = filas_base + (1 if i < extra else 0)
            bloques.append(rows[inicio:inicio + tam])
            inicio += tam

    return bloques, max_por_pagina

def generar_docx_reporte_tutorias_brindadas(
    tutor,
    tutorias,
    plantilla,
    oficio,
    fecha_emision,
    columnas_activas,
    mostrar_col_alumno,
    mostrar_col_fecha,
    mostrar_col_hora,
    mostrar_col_tema,
    mostrar_col_notas,
    tema_dict,
):
    open_plantilla = docx.Document(plantilla.archivo)

    if not open_plantilla.tables:
        return redirect('Reporte-tutorias', pk=tutor.pk)

    _replace_reporte_placeholders(open_plantilla, tutor, oficio, fecha_emision)

    refs = {
        'oficio': _find_reference_paragraph(open_plantilla, f"Oficio No. {oficio}"),
        'fecha': _find_reference_paragraph(open_plantilla, f"Ciudad de México a {fecha_emision}"),
        'asunto': _find_reference_paragraph(open_plantilla, "Asunto Tutorados(as) atendidos"),
    }

    # Ajuste dinámico de filas por página según columnas
    tutorias_validas = _build_report_rows(
    tutorias,
    mostrar_col_alumno,
    mostrar_col_fecha,
    mostrar_col_hora,
    mostrar_col_tema,
    mostrar_col_notas,
    tema_dict,
    )

    bloques, MAX_FILAS = _distribuir_bloques(tutorias_validas, columnas_activas)

    old_table = open_plantilla.tables[0]
    table_style = old_table.style
    parent = old_table._element.getparent()

    placeholder = open_plantilla.add_paragraph()
    parent.insert(parent.index(old_table._element), placeholder._element)
    parent.remove(old_table._element)

    current_element = placeholder._element

    for i, bloque in enumerate(bloques):
        if i > 0:
            page_break = open_plantilla.add_paragraph()
            page_break.add_run().add_break(WD_BREAK.PAGE)
            current_element.addnext(page_break._element)
            current_element = page_break._element
            current_element = _insert_short_header(open_plantilla, current_element, oficio, fecha_emision, refs)

        table = open_plantilla.add_table(rows=1, cols=len(columnas_activas))
        table.style = table_style
        current_element.addnext(table._element)
        current_element = table._element

        header_cells = table.rows[0].cells
        for j, col_name in enumerate(columnas_activas):
            header_cells[j].text = col_name

        for row in bloque:
            row_cells = table.add_row().cells
            for idx, value in enumerate(row):
                row_cells[idx].text = str(value)

        _apply_black_borders(table)

        if len(bloques) > 1 and i == len(bloques) - 1:
            current_element = _push_signature_down(
                open_plantilla,
                current_element,
                len(bloque),
                MAX_FILAS   
            )
    
    _add_page_number_footer(open_plantilla)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename={tutor.last_name}_TUTORIAS_BRINDADAS.docx'
    open_plantilla.save(response)
    return response
